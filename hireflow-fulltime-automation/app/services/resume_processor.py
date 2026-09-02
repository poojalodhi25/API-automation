"""Read candidate resumes from DOCX or PDF files."""

from pathlib import Path

from docx import Document
from pypdf import PdfReader


def extract_resume_text(path: str) -> str:
    """Return plain text from a uploaded resume file."""
    file_path = Path(path)
    if not file_path.exists():
        return ""
    suffix = file_path.suffix.lower()
    if suffix == ".docx":
        document = Document(str(file_path))
        return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text)
    if suffix == ".pdf":
        reader = PdfReader(str(file_path))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages)
    if suffix in {".txt", ".md"}:
        return file_path.read_text(encoding="utf-8", errors="ignore")
    return ""


def write_aligned_resume(original_path: str, aligned_text: str, output_path: str) -> str:
    """Save an aligned resume as a DOCX file."""
    document = Document()
    document.add_heading("Aligned Resume", level=1)
    for block in aligned_text.split("\n"):
        document.add_paragraph(block)
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path
